import docx

def test1():
    doc = docx.Document('test.docx')

    for style in doc.styles: # 目前文件所有格式
        print(style.name)
    print('')

    for i, paragraph in enumerate(doc.paragraphs):
        print('index:', i)
        # print(paragraph.style)
        # print(paragraph.text)
        print(paragraph._element.xml)

    # for table in doc.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             # 打印單元格文本
    #             print(cell.text)
    #             # 打印單元格索引
    #             # print(cell.index)

if __name__ == '__main__':
    test1()